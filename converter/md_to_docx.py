"""
Markdown → docx 转换模块
使用 pypandoc 调用 Pandoc 引擎完成转换，并后处理封面与分页符
"""

import os
import io
import re
import pypandoc
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _check_pandoc():
    """检查 Pandoc 是否可用"""
    try:
        pypandoc.get_pandoc_version()
    except OSError:
        raise RuntimeError(
            "未检测到 Pandoc，请先安装 Pandoc 后重试。\n"
            "Windows: 访问 https://pandoc.org/installing.html 下载安装包\n"
            "Mac: brew install pandoc\n"
            "Linux: sudo apt install pandoc"
        )


def _add_page_break(doc):
    """在文档当前位置插入分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
    return p


def _set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.15):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def _add_cover_page(doc, title="", subtitle="", author="", date_str=""):
    """
    在文档开头插入封面页。
    封面页包含：标题（居中大号）、副标题、作者、日期，底部自动加分页符。
    """
    # 将封面内容插入到文档最前面（位置 0）
    # 先收集现有元素，分离 sectPr（节属性必须保留在最后）
    existing_elements = list(doc.element.body)
    sectPr = None
    for el in existing_elements:
        if el.tag == qn("w:sectPr"):
            sectPr = el
        doc.element.body.remove(el)

    # —— 封面区域 ——
    # 顶部留白（空行）
    for _ in range(6):
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=0, after=0, line_spacing=1.0)

    # 标题
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p, before=12, after=12, line_spacing=1.5)
        run = p.add_run(title)
        run.font.size = Pt(26)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # 副标题
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p, before=6, after=6, line_spacing=1.3)
        run = p.add_run(subtitle)
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # 中间留白
    for _ in range(4):
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=0, after=0, line_spacing=1.5)

    # 作者
    if author:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p, before=6, after=6, line_spacing=1.3)
        run = p.add_run(author)
        run.font.size = Pt(14)

    # 日期
    if date_str:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p, before=6, after=6, line_spacing=1.3)
        run = p.add_run(date_str)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # 封面末尾分页
    _add_page_break(doc)

    # 将原有内容追加回来（sectPr 必须放在最后）
    for el in existing_elements:
        if el.tag != qn("w:sectPr"):
            doc.element.body.append(el)
    if sectPr is not None:
        doc.element.body.append(sectPr)


def _extract_cover_info(md_text: str) -> dict:
    """
    从 Markdown 文本中提取封面信息（标题、副标题、作者、日期）。
    覆盖三种常见模式：
    1. # 标题 / ## 副标题 / **作者：** xxx / **日期：** xxx
    2. 开头连续元数据行
    """
    info = {"title": "", "subtitle": "", "author": "", "date_str": ""}

    lines = md_text.strip().split("\n")

    # 提取第一个一级标题作为封面标题
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith("# ") and not stripped.startswith("## "):
            info["title"] = stripped[2:].strip()
            break
        elif stripped.startswith("#"):
            info["title"] = stripped.lstrip("#").strip()
            break

    # 提取第二个一级标题或第一个二级标题作为副标题
    h1_count = 0
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# ") and not stripped.startswith("## "):
            h1_count += 1
            if h1_count == 2:
                info["subtitle"] = stripped[2:].strip()
                break
        elif h1_count == 1 and stripped.startswith("## "):
            if not info["subtitle"]:
                info["subtitle"] = stripped[3:].strip()
                break

    # 提取作者和日期（匹配 **作者：** xxx 或 **日期：** xxx 等）
    for line in lines[:20]:  # 仅在前 20 行搜索
        stripped = line.strip()
        author_match = re.match(r"\*\*(作者|Author|姓名)[：:]\s*\*\*\s*(.+)", stripped, re.IGNORECASE)
        date_match = re.match(r"\*\*(日期|Date|时间)[：:]\s*\*\*\s*(.+)", stripped, re.IGNORECASE)

        if author_match and not info["author"]:
            info["author"] = author_match.group(2).strip()
        if date_match and not info["date_str"]:
            info["date_str"] = date_match.group(2).strip()

    return info


def _parse_metadata_from_md(md_text: str) -> dict:
    """从 Markdown 中提取基础元数据（作者、日期）。兼容 _extract_cover_info 未覆盖的情况。"""
    info = _extract_cover_info(md_text)
    lines = md_text.strip().split("\n")

    # 尝试匹配纯文本行中的作者/日期信息
    for line in lines[:20]:
        stripped = line.strip()
        # 匹配 "作者：xxx" 或 "Author: xxx"
        if not info["author"]:
            m = re.match(r"(?:作者|Author|姓名)[：:]\s*(.+)", stripped, re.IGNORECASE)
            if m:
                info["author"] = m.group(1).strip().rstrip("*").lstrip("*").strip()
        # 匹配 "日期：xxx" 或 "Date: xxx"
        if not info["date_str"]:
            m = re.match(r"(?:日期|Date|时间)[：:]\s*(.+)", stripped, re.IGNORECASE)
            if m:
                info["date_str"] = m.group(1).strip().rstrip("*").lstrip("*").strip()

    return info


def post_process_docx(docx_bytes: bytes, md_source: str = "") -> bytes:
    """
    对 Pandoc 生成的 docx 进行后处理：
    1. 根据 Markdown 内容自动提取封面信息并插入封面页
    2. 确保所有 \\newpage 标记已由 Pandoc 转为分页符（Pandoc 原生支持）
    3. 统一正文字体与段落间距

    参数：
        docx_bytes: Pandoc 生成的原始 docx 二进制数据
        md_source: 原始 Markdown 文本（用于提取封面信息）

    返回：
        bytes: 后处理后的 docx 二进制数据
    """
    try:
        doc = Document(io.BytesIO(docx_bytes))
    except Exception:
        # 如果 python-docx 无法解析，直接返回原始数据
        return docx_bytes

    # —— 封面页处理 ——
    if md_source:
        cover = _parse_metadata_from_md(md_source)
        # 至少要有标题才生成封面
        if cover["title"]:
            _add_cover_page(
                doc,
                title=cover["title"],
                subtitle=cover["subtitle"],
                author=cover.get("author", ""),
                date_str=cover.get("date_str", ""),
            )

    # —— 全局段落样式规范化 ——
    try:
        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        style.font.name = "微软雅黑"
        if style.element.rPr is not None:
            style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    except Exception:
        pass  # 样式不存在或无法修改时跳过

    # —— 标题样式规范化 ——
    for heading_style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        hs = doc.styles.get(heading_style_name, None)
        if hs:
            if heading_style_name == "Heading 1":
                hs.font.size = Pt(18)
                hs.font.bold = True
            elif heading_style_name == "Heading 2":
                hs.font.size = Pt(15)
                hs.font.bold = True
            elif heading_style_name == "Heading 3":
                hs.font.size = Pt(13)
                hs.font.bold = True

    # 保存到 bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def convert_md_to_docx(md_text: str, template_path: str = None, enable_post_process: bool = True) -> bytes:
    """
    将 Markdown 文本转换为 docx 格式的二进制数据

    参数：
        md_text (str): 待转换的 Markdown 文本内容
        template_path (str, optional): Pandoc reference.docx 模板路径
        enable_post_process (bool): 是否启用封面/分页后处理，默认 True

    返回：
        bytes: docx 文件的二进制数据，供下载使用

    异常：
        RuntimeError: Pandoc 未安装时抛出
        Exception: 转换过程发生错误时抛出
    """
    _check_pandoc()

    try:
        extra_args = []

        # 若提供了模板文件且存在，则使用 reference-doc 参数
        if template_path and os.path.isfile(template_path):
            extra_args.extend(["--reference-doc", template_path])

        # 调用 pypandoc 进行转换（Pandoc 原生支持 \newpage 分页符）
        docx_bytes = pypandoc.convert_text(
            md_text,
            "docx",
            format="markdown",
            extra_args=extra_args,
        )

        # 后处理：封面页 + 样式规范化
        if enable_post_process:
            docx_bytes = post_process_docx(docx_bytes, md_source=md_text)

        return docx_bytes

    except Exception as e:
        raise Exception(f"Markdown 转 docx 失败：{str(e)}")
